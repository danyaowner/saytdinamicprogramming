#!/usr/bin/env python3
"""
Validate the generated coursework.docx - check structure, styles, TOC, formatting.
"""
import sys
import os
import re
try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed")
    sys.exit(1)

doc_path = "coursework.docx"
if not os.path.exists(doc_path):
    alt_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), doc_path)
    if os.path.exists(alt_path):
        doc_path = alt_path
    else:
        print(f"File not found: {doc_path}")
        # try /tmp
        doc_path = "/tmp/coursework.docx"
        if not os.path.exists(doc_path):
            print(f"Also not found: {doc_path}")
            sys.exit(1)

print(f"Checking: {doc_path}")
print("=" * 60)

doc = Document(doc_path)
errors = []
warnings = []
passes = []

# 1. Stats
print("\nSTATS")
print("-" * 40)
total_pars = len(doc.paragraphs)
total_tables = len(doc.tables)
print(f"  Paragraphs: {total_pars}")
print(f"  Tables: {total_tables}")

sec = doc.sections[0]
pw_cm = sec.page_width / 360000
ph_cm = sec.page_height / 360000
print(f"  Page: {pw_cm:.1f}x{ph_cm:.1f} cm")
print(f"  Margins: L={sec.left_margin/360000:.1f} R={sec.right_margin/360000:.1f} T={sec.top_margin/360000:.1f} B={sec.bottom_margin/360000:.1f} cm")

if abs(pw_cm - 21.0) < 0.5 and abs(ph_cm - 29.7) < 0.5:
    passes.append("Page size A4 (21x29.7cm)")
if abs(sec.left_margin/360000 - 3.0) < 0.3:
    passes.append("Left margin ~3cm")
if abs(sec.right_margin/360000 - 1.5) < 0.3:
    passes.append("Right margin ~1.5cm")

# 2. Normal style
normal = doc.styles['Normal']
print(f"\nNormal font: {normal.font.name}, size: {normal.font.size}")
if normal.font.name == 'Times New Roman':
    passes.append("Normal: Times New Roman")
else:
    errors.append(f"Normal font is {normal.font.name}")

# 4. Headings
print("\nHEADINGS")
print("-" * 40)
h1_found = []
h2_found = []

for p in doc.paragraphs:
    s = p.style.name
    t = p.text.strip()
    if s == 'Heading 1' and t:
        h1_found.append(t[:80])
    elif s == 'Heading 2' and t:
        h2_found.append(t[:80])

print(f"Heading 1 found: {len(h1_found)}")
for h in h1_found:
    print(f"  - {h}")
print(f"Heading 2 found: {len(h2_found)}")
for h in h2_found:
    print(f"  - {h}")

expected_h1 = ["ВВЕДЕНИЕ", "ТЕОРЕТИЧЕСКИЕ", "МАТЕМАТИЧЕСКАЯ", "ПРОГРАММНАЯ", "ЗАКЛЮЧЕНИЕ", "СПИСОК", "ПРИЛОЖЕНИЕ"]
for exp in expected_h1:
    if any(exp in h.upper() for h in h1_found):
        passes.append(f"Section '{exp}' has Heading 1")
    else:
        errors.append(f"Section '{exp}' MISSING Heading 1")

# 5. TOC check
toc_found = False
for p in doc.paragraphs:
    for r in p._p.findall(qn('w:r')):
        for instr in r.findall(qn('w:instrText')):
            if instr.text and 'TOC' in instr.text:
                toc_found = True
                print(f"\nTOC field: {instr.text.strip()}")

if toc_found:
    passes.append("TOC field present")
else:
    errors.append("TOC field NOT found")

# Check OGLABLENIE heading
for p in doc.paragraphs:
    if p.text.strip() == 'ОГЛАВЛЕНИЕ':
        print(f"TOC title style: {p.style.name}")
        if p.style.name == 'Heading 1':
            warnings.append("'ОГЛАВЛЕНИЕ' has Heading 1 style!")
        else:
            passes.append("'ОГЛАВЛЕНИЕ' does NOT have Heading 1 (good)")

# 6. Check references
print("\nREFERENCES")
print("-" * 40)
all_text = ' '.join(p.text for p in doc.paragraphs)
refs = set()
for m in re.finditer(r'\[(\d+)\]', all_text):
    refs.add(int(m.group(1)))
print(f"References used: {sorted(refs)}")
if refs:
    passes.append(f"References found: [{','.join(map(str,sorted(refs)))}]")

# 7. Summary
print("\n" + "=" * 60)
print("RESULTS")
print("=" * 60)
print(f"\nPASSES: {len(passes)}")
for p in passes:
    print(f"  [OK] {p}")
if warnings:
    print(f"\nWARNINGS: {len(warnings)}")
    for w in warnings:
        print(f"  [!] {w}")
if errors:
    print(f"\nERRORS: {len(errors)}")
    for e in errors:
        print(f"  [FAIL] {e}")
else:
    print("\nALL CHECKS PASSED!")

print("\nWhat to do in Word:")
print("1. Open coursework.docx")
print("2. Right-click TOC -> 'Update Field' (or press F9)")
print("3. Choose 'Update entire table'")

sys.exit(0 if not errors else 1)
